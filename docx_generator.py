"""
docx_generator.py — Generate DOCX output matching AIIMS Patna NCS-EMG report format.

Font spec (from reference PDF):
  - "Department of Physiology" / "AIIMS Patna": Times New Roman 14pt bold+underlined
  - Demographic table: Segoe UI Historic 10pt
  - NCS/EMG/F-Wave/VEP tables + section headings: Segoe UI Historic 8pt
  - Report title (NCS SUMMARY REPORT): Segoe UI Historic 14pt bold+italic+underlined
  - Summary:, Impression: labels: Calibri 11pt bold+underlined
  - Summary subheadings (Motor NCS:, etc.): Calibri 11pt bold
  - Body text (summary content, impression, signature): Calibri 11pt
"""

import os
import io
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def _find_libreoffice():
    candidates = [
        '/opt/homebrew/bin/soffice',
        '/usr/local/bin/soffice',
        '/Applications/LibreOffice.app/Contents/MacOS/soffice',
    ]
    for p in candidates:
        if os.path.exists(p):
            return p
    return None

LIBREOFFICE = _find_libreoffice()

FONT_TIMES   = 'Times New Roman'
FONT_SEGOE   = 'Segoe UI Historic'
FONT_CALIBRI = 'Calibri'
LOGO_PATH = os.path.join(os.path.dirname(__file__), 'static', 'aiims_patna_logo.png')


# ── XML helpers ───────────────────────────────────────────────────────────────

def _set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        attrs = kwargs.get(edge, {})
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), attrs.get('val', 'single'))
        tag.set(qn('w:sz'), attrs.get('sz', '4'))
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), attrs.get('color', '000000'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def _thin():
    return {'val': 'single', 'sz': '4', 'color': '000000'}


def _no_border():
    return {'val': 'none', 'sz': '0', 'color': 'FFFFFF'}


def _set_cell_no_shading(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFFFFF')
    tcPr.append(shd)


def _set_para_spacing(para, before=0, after=0):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))
    pPr.append(spacing)


def _run(para, text, bold=False, italic=False, underline=False, size=10, font=None):
    run = para.add_run(text)
    run.font.name = font if font else FONT_TIMES
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    return run


def _set_table_cell_margins(table, top=40, left=40, bottom=40, right=40):
    """Set uniform cell margins for the entire table (in dxa = twips, ~0.7mm each side)."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblCellMar = OxmlElement('w:tblCellMar')
    for side, val in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tblCellMar.append(el)
    tblPr.append(tblCellMar)


# ── Demographics table ────────────────────────────────────────────────────────

def _build_demo_table(doc, patient):
    fields_left = [
        ("Patient's Name : ", patient.get('name', '')),
        ("ID: ",              patient.get('id', '')),
        ("Referring Department: ", patient.get('referring_dept', '')),
        ("Diagnosis: ",       patient.get('diagnosis', '')),
    ]
    age = patient.get('age', '')
    sex = patient.get('sex', '')
    age_sex = f"{age} years /{sex}" if age or sex else ''

    fields_right = [
        ("Age/Sex: ",      age_sex),
        ("Date: ",         patient.get('date', '')),
        ("Doctor: ",       patient.get('doctor', '')),
        ("Performed by: ", patient.get('performed_by', 'Mr. Manish Kumar')),
    ]

    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for (lbl_l, val_l), (lbl_r, val_r) in zip(fields_left, fields_right):
        row = table.add_row()
        for cell, lbl, val in [(row.cells[0], lbl_l, val_l),
                               (row.cells[1], lbl_r, val_r)]:
            _set_cell_no_shading(cell)
            _set_cell_border(cell, top=_thin(), left=_thin(),
                             bottom=_thin(), right=_thin())
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _set_para_spacing(p, before=20, after=20)
            _run(p, lbl, bold=True, size=10, font=FONT_SEGOE)
            _run(p, val, bold=False, size=10, font=FONT_SEGOE)

    return table


# ── NCS data tables ───────────────────────────────────────────────────────────

def _add_section_heading(doc, text, size=10):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(p, before=60, after=20)
    _run(p, text, bold=True, size=size, font=FONT_SEGOE)
    return p


def _build_ncs_table(doc, rows, header_cols, section_title, value_key_map):
    if not rows:
        return

    _add_section_heading(doc, section_title)

    n_cols = len(header_cols)
    table = doc.add_table(rows=0, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    _set_table_cell_margins(table, top=30, left=40, bottom=30, right=40)

    # Header row
    hdr_row = table.add_row()
    for cell, hdr in zip(hdr_row.cells, header_cols):
        _set_cell_no_shading(cell)
        _set_cell_border(cell, top=_thin(), left=_thin(),
                         bottom=_thin(), right=_thin())
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_spacing(p, before=20, after=20)
        _run(p, hdr, bold=True, size=10, font=FONT_SEGOE)

    # Data rows
    current_group = None
    for row in rows:
        group = row.get('group', '')
        if group != current_group:
            current_group = group
            grp_row = table.add_row()
            merged = grp_row.cells[0].merge(grp_row.cells[-1])
            _set_cell_no_shading(merged)
            _set_cell_border(merged, top=_thin(), left=_thin(),
                             bottom=_thin(), right=_thin())
            p = merged.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _set_para_spacing(p, before=20, after=20)
            _run(p, group, bold=True, size=10, font=FONT_SEGOE)

        data_row = table.add_row()
        for i, key in enumerate(value_key_map):
            cell = data_row.cells[i]
            val = row.get(key)
            text = str(val) if val is not None else ''
            _set_cell_no_shading(cell)
            _set_cell_border(cell, top=_thin(), left=_thin(),
                             bottom=_thin(), right=_thin())
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_para_spacing(p, before=15, after=15)
            _run(p, text, size=10, font=FONT_SEGOE)


def _build_fwave_table(doc, fwave_rows):
    if not fwave_rows:
        return

    _add_section_heading(doc, 'F-Wave Findings')

    headers = ['Test', 'Fmin lat., ms', 'M lat., ms', 'Fmin-M lat., ms', 'Max Vprox, m/s']
    keys    = ['test_no', 'fmin_lat_ms', 'm_lat_ms', 'fmin_m_lat_ms', 'max_vprox']

    table = doc.add_table(rows=0, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    _set_table_cell_margins(table, top=30, left=40, bottom=30, right=40)

    hdr_row = table.add_row()
    for cell, hdr in zip(hdr_row.cells, headers):
        _set_cell_no_shading(cell)
        _set_cell_border(cell, top=_thin(), left=_thin(),
                         bottom=_thin(), right=_thin())
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_spacing(p, before=20, after=20)
        _run(p, hdr, bold=True, size=10, font=FONT_SEGOE)

    current_group = None
    for row in fwave_rows:
        group = row.get('group', '')
        if group != current_group:
            current_group = group
            grp_row = table.add_row()
            merged = grp_row.cells[0].merge(grp_row.cells[-1])
            _set_cell_no_shading(merged)
            _set_cell_border(merged, top=_thin(), left=_thin(),
                             bottom=_thin(), right=_thin())
            p = merged.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _set_para_spacing(p, before=20, after=20)
            _run(p, group, bold=True, size=10, font=FONT_SEGOE)

        data_row = table.add_row()
        for key in keys:
            cell = data_row.cells[keys.index(key)]
            val = row.get(key)
            _set_cell_no_shading(cell)
            _set_cell_border(cell, top=_thin(), left=_thin(),
                             bottom=_thin(), right=_thin())
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_para_spacing(p, before=15, after=15)
            _run(p, str(val) if val is not None else '', size=10, font=FONT_SEGOE)


def _set_vmerge(cell, restart=False):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vmerge = OxmlElement('w:vMerge')
    if restart:
        vmerge.set(qn('w:val'), 'restart')
    tcPr.append(vmerge)


def _set_vAlign_center(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'center')
    tcPr.append(vAlign)


def _header_cell(cell, text, bold=True, size=10):
    _set_cell_no_shading(cell)
    _set_cell_border(cell, top=_thin(), left=_thin(), bottom=_thin(), right=_thin())
    _set_vAlign_center(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(p, before=20, after=20)
    _run(p, text, bold=bold, size=size, font=FONT_SEGOE)


def _build_emg_table(doc, emg_form_data):
    if not emg_form_data:
        return

    _add_section_heading(doc, 'EMG Findings')

    any_notes = any(bool((row.get('notes') or '').strip()) for row in emg_form_data if row.get('muscle'))

    keys = ['muscle', 'insertion', 'resting', 'amplitude', 'duration',
            'polyphasic', 'recruitment', 'interference']
    if any_notes:
        keys.append('notes')

    num_cols = 9 if any_notes else 8
    table = doc.add_table(rows=0, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    _set_table_cell_margins(table, top=30, left=40, bottom=30, right=40)

    # Header row 1
    hdr1 = table.add_row()
    vmerge_hdrs = [(0, 'Muscle'), (1, 'Insertion\nEMG'), (2, 'Resting\nEMG'), (7, 'Interference')]
    if any_notes:
        vmerge_hdrs.append((8, 'Notes'))
    for col, text in vmerge_hdrs:
        _header_cell(hdr1.cells[col], text)
        _set_vmerge(hdr1.cells[col], restart=True)
    mup_cell = hdr1.cells[3].merge(hdr1.cells[6])
    _header_cell(mup_cell, 'MUP Findings')

    # Header row 2
    hdr2 = table.add_row()
    vmerge_cont = [0, 1, 2, 7] + ([8] if any_notes else [])
    for col in vmerge_cont:
        cell = hdr2.cells[col]
        _set_cell_no_shading(cell)
        _set_cell_border(cell, top=_thin(), left=_thin(), bottom=_thin(), right=_thin())
        _set_vmerge(cell)
    for col, text in [(3, 'Amplitude'), (4, 'Duration'), (5, 'Polyphasic'), (6, 'Recruitment')]:
        _header_cell(hdr2.cells[col], text)

    MUSCLE_KEY_UL = (
        'DL=Deltoid, SS=Supraspinatus, IS=Infraspinatus, BB=Biceps Brachii, '
        'TR=Triceps, BR=Brachioradialis, FCR=Flexor Carpi Radialis, '
        'FCU=Flexor Carpi Ulnaris, APB=Abductor Pollicis Brevis, '
        'ADM=Abductor Digiti Minimi, IO=First Dorsal Interosseous, '
        'LB=Lumbrical Brevis, ED=Extensor Digitorum'
    )
    MUSCLE_KEY_LL = (
        'VM=Vastus Medialis, VL=Vastus Lateralis, GN=Gastrocnemius, '
        'FDL=Flexor Digitorum Longus, FHL=Flexor Hallucis Longus, '
        'TA=Tibialis Anterior, EDL=Extensor Digitorum Longus, '
        'EDB=Extensor Digitorum Brevis, AHL=Abductor Hallucis, HS=Hamstrings'
    )

    for row in emg_form_data:
        if not row.get('muscle'):
            continue
        side_raw = row.get('side', '')
        sides = ['L', 'R'] if side_raw == 'B' else [side_raw]

        for s in sides:
            data_row = table.add_row()
            cant_recruit = (row.get('recruitment') == "Patient couldn't recruit")

            if cant_recruit:
                for i, key in enumerate(['muscle', 'insertion', 'resting']):
                    cell = data_row.cells[i]
                    val = row.get(key, '') or ''
                    if key == 'muscle':
                        val = f"{s} {val}" if s else val
                    _set_cell_no_shading(cell)
                    _set_cell_border(cell, top=_thin(), left=_thin(),
                                     bottom=_thin(), right=_thin())
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _set_para_spacing(p, before=15, after=15)
                    _run(p, str(val), size=10, font=FONT_SEGOE)
                merged = data_row.cells[3].merge(data_row.cells[7])
                _set_cell_no_shading(merged)
                _set_cell_border(merged, top=_thin(), left=_thin(),
                                 bottom=_thin(), right=_thin())
                p = merged.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _set_para_spacing(p, before=15, after=15)
                _run(p, "Patient could not recruit", size=10, font=FONT_SEGOE)
                if any_notes:
                    cell = data_row.cells[8]
                    _set_cell_no_shading(cell)
                    _set_cell_border(cell, top=_thin(), left=_thin(),
                                     bottom=_thin(), right=_thin())
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _set_para_spacing(p, before=15, after=15)
                    _run(p, row.get('notes', '') or '', size=10, font=FONT_SEGOE)
            else:
                for i, key in enumerate(keys):
                    cell = data_row.cells[i]
                    val = row.get(key, '') or ''
                    if key == 'muscle':
                        val = f"{s} {val}" if s else val
                    _set_cell_no_shading(cell)
                    _set_cell_border(cell, top=_thin(), left=_thin(),
                                     bottom=_thin(), right=_thin())
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _set_para_spacing(p, before=15, after=15)
                    _run(p, str(val), size=10, font=FONT_SEGOE)

    key_p1 = doc.add_paragraph()
    _set_para_spacing(key_p1, before=60, after=0)
    _run(key_p1, 'Upper Limb: ', bold=True, size=7, font=FONT_SEGOE)
    _run(key_p1, MUSCLE_KEY_UL, size=7, font=FONT_SEGOE)

    key_p2 = doc.add_paragraph()
    _set_para_spacing(key_p2, before=0, after=60)
    _run(key_p2, 'Lower Limb: ', bold=True, size=7, font=FONT_SEGOE)
    _run(key_p2, MUSCLE_KEY_LL, size=7, font=FONT_SEGOE)


def _build_vep_table(doc, vep_rows):
    if not vep_rows:
        return

    _add_section_heading(doc, 'Visual Evoked Potential (VEP)')

    headers = ['N', 'Rec sites', 'N75 lat', 'P100 lat', 'N145 lat',
               'P100-N145 ampl', 'Stim side', 'Stimulus', 'Stim dur']
    keys    = ['n', 'rec_sites', 'n75_lat', 'p100_lat', 'n145_lat',
               'p100_n145_ampl', 'stim_side', 'stimulus', 'stim_dur']

    table = doc.add_table(rows=0, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    _set_table_cell_margins(table, top=30, left=40, bottom=30, right=40)

    hdr_row = table.add_row()
    for cell, hdr in zip(hdr_row.cells, headers):
        _set_cell_no_shading(cell)
        _set_cell_border(cell, top=_thin(), left=_thin(),
                         bottom=_thin(), right=_thin())
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_spacing(p, before=20, after=20)
        _run(p, hdr, bold=True, size=10, font=FONT_SEGOE)

    for row in vep_rows:
        data_row = table.add_row()
        for i, key in enumerate(keys):
            cell = data_row.cells[i]
            val = row.get(key)
            _set_cell_no_shading(cell)
            _set_cell_border(cell, top=_thin(), left=_thin(),
                             bottom=_thin(), right=_thin())
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_para_spacing(p, before=15, after=15)
            _run(p, str(val) if val is not None else '', size=10, font=FONT_SEGOE)


# ── Waveform helpers ──────────────────────────────────────────────────────────

def _autocrop_png(png_path):
    try:
        from PIL import Image, ImageChops
        img = Image.open(png_path).convert('RGB')
        bg = Image.new('RGB', img.size, (255, 255, 255))
        diff = ImageChops.difference(img, bg)
        bbox = diff.getbbox()
        if bbox:
            pad = 10
            w, h = img.size
            x0 = max(0, bbox[0] - pad)
            y0 = max(0, bbox[1] - pad)
            x1 = min(w, bbox[2] + pad)
            y1 = min(h, bbox[3] + pad)
            img.crop((x0, y0, x1, y1)).save(png_path)
    except Exception:
        pass


def _build_emg_waveforms(doc, emg_images):
    if not emg_images:
        return
    available = [(img['label'], img['png_bytes'])
                 for img in emg_images if img.get('png_bytes')]
    if not available:
        return

    _add_section_heading(doc, 'Waveform Tracings')
    MAX_WIDTH_CM = 13.0

    for label, png_bytes in available:
        img_width_arg = Cm(MAX_WIDTH_CM)
        try:
            from PIL import Image as _PILImage
            import io as _io
            with _PILImage.open(_io.BytesIO(png_bytes)) as _im:
                px_w, _ = _im.size
                dpi = _im.info.get('dpi', (96, 96))
                dpi_x = dpi[0] if isinstance(dpi, (tuple, list)) else dpi
                nat_w_cm = px_w / dpi_x * 2.54
                if nat_w_cm <= MAX_WIDTH_CM:
                    img_width_arg = Cm(nat_w_cm)
        except Exception:
            pass

        lp = doc.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp.paragraph_format.keep_with_next = True
        _set_para_spacing(lp, before=50, after=4)
        _run(lp, label, bold=True, size=9, font=FONT_SEGOE)

        ip = doc.add_paragraph()
        ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_spacing(ip, before=0, after=30)
        import io as _io
        ip.add_run().add_picture(_io.BytesIO(png_bytes), width=img_width_arg)


def _build_ncs_waveforms(doc, ncs_images):
    if not ncs_images:
        return
    available = [(img['label'], img['png_bytes'])
                 for img in ncs_images if img.get('png_bytes')]
    if not available:
        return

    # No section heading — waveforms follow directly after tables (matches reference)
    MAX_WIDTH_CM = 13.0

    for label, png_bytes in available:
        img_width_arg = Cm(MAX_WIDTH_CM)
        try:
            from PIL import Image as _PILImage
            import io as _io
            with _PILImage.open(_io.BytesIO(png_bytes)) as _im:
                px_w, _ = _im.size
                dpi = _im.info.get('dpi', (96, 96))
                dpi_x = dpi[0] if isinstance(dpi, (tuple, list)) else dpi
                nat_w_cm = px_w / dpi_x * 2.54
                if nat_w_cm <= MAX_WIDTH_CM:
                    img_width_arg = Cm(nat_w_cm)
        except Exception:
            pass

        lp = doc.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp.paragraph_format.keep_with_next = True
        _set_para_spacing(lp, before=50, after=4)
        _run(lp, label, bold=True, size=9, font=FONT_SEGOE)

        ip = doc.add_paragraph()
        ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_spacing(ip, before=0, after=30)
        import io as _io
        ip.add_run().add_picture(_io.BytesIO(png_bytes), width=img_width_arg)


# ── Main entry point ──────────────────────────────────────────────────────────

def generate_docx(parsed_data, report_dict, emg_form_data, report_types,
                  lab_name='NCS-EMG LAB, DEPARTMENT OF PHYSIOLOGY, AIIMS PATNA',
                  emg_images=None, ncs_images=None, include_ncs_tables=True, include_vep_table=True):
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    # ── Logo ──────────────────────────────────────────────────────────────────
    if os.path.exists(LOGO_PATH):
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_spacing(logo_para, before=0, after=60)
        logo_para.add_run().add_picture(LOGO_PATH, width=Inches(1.1))

    # ── Institution header — Times New Roman 14pt bold+underlined ─────────────
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(p1, before=0, after=0)
    _run(p1, 'Department of Physiology', bold=True, underline=True, size=14, font=FONT_TIMES)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(p2, before=0, after=80)
    _run(p2, 'AIIMS Patna', bold=True, underline=True, size=14, font=FONT_TIMES)

    # ── Demographics table — Segoe UI Historic 10pt ───────────────────────────
    patient = parsed_data.get('patient', {})
    _build_demo_table(doc, patient)

    # ── Report title — Segoe UI Historic 14pt bold+italic+underlined ──────────
    doc.add_paragraph()  # spacer

    if 'NCS' in report_types and 'EMG' in report_types:
        title = 'NCS-EMG REPORT'
    elif 'VEP' in report_types:
        title = 'VEP REPORT'
    elif 'EMG' in report_types:
        title = 'EMG REPORT'
    else:
        title = 'NCS SUMMARY REPORT'

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(t, before=80, after=80)
    _run(t, title, bold=True, italic=True, underline=True, size=14, font=FONT_SEGOE)

    # ── NCS Tables — Segoe UI Historic 8pt, minimal cell padding ─────────────
    if 'NCS' in report_types and include_ncs_tables:
        # Motor CV: 11 columns (no Vel. norm / Vel. dev)
        motor_headers = ['Test', 'Stimulation site', 'Lat., ms', 'Ampl., mV',
                         'Dur., ms', 'Area, mV×ms', 'Stim., mA', 'Stim., ms',
                         'Dist., mm', 'Time, ms', 'Vel., m/s']
        motor_keys    = ['test_no', 'site', 'lat_ms', 'ampl_mv', 'dur_ms', 'area',
                         'stim_ma', 'stim_ms', 'dist_mm', 'time_ms', 'vel_ms']
        _build_ncs_table(doc, parsed_data.get('motor_ncs', []),
                         motor_headers, 'Motor CV', motor_keys)

        # Sensory CV: 13 columns (with Vel. norm / Vel. dev)
        sensory_headers = ['Test', 'Site', 'Lat., ms', 'Ampl., µV', 'Dur., ms',
                           'Area, nV×s', 'Stim., mA', 'Stim., ms', 'Dist., mm',
                           'Time, ms', 'Vel., m/s', 'Vel. norm, m/s', 'Vel. dev., %']
        sensory_keys    = ['test_no', 'site', 'lat_ms', 'ampl_uv', 'dur_ms', 'area',
                           'stim_ma', 'stim_ms', 'dist_mm', 'time_ms', 'vel_ms',
                           'vel_norm', 'vel_dev_pct']
        _build_ncs_table(doc, parsed_data.get('sensory_ncs', []),
                         sensory_headers, 'Sensory CV', sensory_keys)

        _build_fwave_table(doc, parsed_data.get('f_waves', []))

    # ── EMG Table ─────────────────────────────────────────────────────────────
    if 'EMG' in report_types and emg_form_data:
        _build_emg_table(doc, emg_form_data)

    # ── NCS Waveforms (no section heading — matches reference) ────────────────
    if ncs_images:
        _build_ncs_waveforms(doc, ncs_images)

    # ── EMG Waveform Tracings ─────────────────────────────────────────────────
    if 'EMG' in report_types and emg_images:
        _build_emg_waveforms(doc, emg_images)

    # ── VEP Table ─────────────────────────────────────────────────────────────
    if 'VEP' in report_types and parsed_data.get('vep') and include_vep_table:
        _build_vep_table(doc, parsed_data['vep'])

    # ── Summary — Calibri 11pt ────────────────────────────────────────────────
    doc.add_paragraph()  # spacer

    s_label = doc.add_paragraph()
    _set_para_spacing(s_label, before=0, after=20)
    _run(s_label, 'Summary:', bold=True, underline=True, size=11, font=FONT_CALIBRI)

    def _add_section(heading, content):
        if not content:
            return
        h_p = doc.add_paragraph()
        _set_para_spacing(h_p, before=40, after=0)
        _run(h_p, heading, bold=True, size=11, font=FONT_CALIBRI)
        c_p = doc.add_paragraph()
        _set_para_spacing(c_p, before=0, after=20)
        _run(c_p, content, size=11, font=FONT_CALIBRI)

    intro = report_dict.get('intro', '')
    if intro:
        p = doc.add_paragraph()
        _set_para_spacing(p, before=0, after=40)
        _run(p, intro, size=11, font=FONT_CALIBRI)

    if 'NCS' in report_types:
        _add_section('Motor NCS:', report_dict.get('motor_summary', ''))
        _add_section('Sensory NCS:', report_dict.get('sensory_summary', ''))
        if report_dict.get('fwave_summary'):
            _add_section('F-Waves:', report_dict.get('fwave_summary', ''))

    if 'EMG' in report_types and report_dict.get('emg_summary'):
        _add_section('EMG:', report_dict.get('emg_summary', ''))

    if 'VEP' in report_types and report_dict.get('vep_summary'):
        _add_section('VEP:', report_dict.get('vep_summary', ''))

    # ── Impression — Calibri 11pt bold+underlined ─────────────────────────────
    imp_label = doc.add_paragraph()
    _set_para_spacing(imp_label, before=60, after=20)
    _run(imp_label, 'Impression:', bold=True, underline=True, size=11, font=FONT_CALIBRI)

    for line in report_dict.get('impression', '').split('\n'):
        lp = doc.add_paragraph()
        _set_para_spacing(lp, before=0, after=20)
        _run(lp, line, size=11, font=FONT_CALIBRI)

    # ── Signature — Calibri 11pt ──────────────────────────────────────────────
    sig_p = doc.add_paragraph()
    sig_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_para_spacing(sig_p, before=120, after=0)
    _run(sig_p, 'Senior Resident/Consultant', bold=True, size=11, font=FONT_CALIBRI)

    dept_p = doc.add_paragraph()
    dept_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_para_spacing(dept_p, before=0, after=0)
    _run(dept_p, 'Dept. of Physiology', size=11, font=FONT_CALIBRI)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
